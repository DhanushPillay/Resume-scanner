import re
import pdfplumber
import docx
import spacy
from dateutil import parser as date_parser
from datetime import datetime
from dateutil.relativedelta import relativedelta

# Lazy load Spacy model (to prevent import-time failures on deployment)
_nlp = None

def get_nlp():
    """Lazy load spaCy model on first use."""
    global _nlp
    if _nlp is None:
        try:
            _nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("Spacy model 'en_core_web_sm' not found. Run: python -m spacy download en_core_web_sm")
            _nlp = False  # Mark as failed so we don't retry
    return _nlp if _nlp else None

# Common job titles for better detection
JOB_TITLES = [
    "software engineer", "senior software engineer", "staff engineer", "principal engineer",
    "frontend developer", "backend developer", "full stack developer", "fullstack developer",
    "data scientist", "data analyst", "machine learning engineer", "ml engineer", "ai engineer",
    "devops engineer", "sre", "site reliability engineer", "cloud engineer", "cloud architect",
    "product manager", "project manager", "engineering manager", "technical lead", "tech lead",
    "cto", "ceo", "cfo", "vp of engineering", "director of engineering", "head of engineering",
    "intern", "software intern", "developer", "programmer", "analyst", "consultant",
    "architect", "solution architect", "systems architect", "software architect",
    "qa engineer", "test engineer", "automation engineer", "quality assurance",
    "ui/ux designer", "ux designer", "ui designer", "product designer", "graphic designer",
    "database administrator", "dba", "system administrator", "network engineer",
    # Additional modern titles
    "platform engineer", "security engineer", "devsecops engineer", "data engineer",
    "mlops engineer", "prompt engineer", "ai researcher", "blockchain developer",
    "mobile developer", "ios developer", "android developer", "flutter developer",
    "frontend engineer", "backend engineer", "reliability engineer"
]

# Expanded tech skills for extraction (including modern tech)
TECH_SKILLS = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "go", "golang", "rust", "ruby",
    "php", "swift", "kotlin", "scala", "r", "matlab", "perl", "shell", "bash", "powershell",
    "objective-c", "dart", "elixir", "clojure", "haskell", "lua", "groovy", "f#",
    
    # Frontend
    "html", "css", "sass", "less", "tailwind", "tailwindcss", "bootstrap", "material ui",
    "react", "react.js", "reactjs", "angular", "vue", "vue.js", "vuejs", "svelte", 
    "next.js", "nextjs", "nuxt", "nuxt.js", "gatsby", "remix", "astro", "solid.js",
    "jquery", "webpack", "vite", "rollup", "parcel", "esbuild",
    "shadcn", "radix ui", "chakra ui", "ant design", "styled-components",
    
    # Backend Frameworks
    "node.js", "nodejs", "express", "express.js", "fastapi", "flask", "django",
    "spring", "spring boot", "rails", "ruby on rails", "laravel", "symfony",
    "asp.net", ".net", "dotnet", "gin", "echo", "fiber", "nestjs", "nest.js",
    "fastify", "koa", "hapi", "phoenix", "actix",
    
    # Cloud & Infrastructure
    "aws", "amazon web services", "azure", "gcp", "google cloud", "google cloud platform",
    "heroku", "vercel", "netlify", "digitalocean", "linode", "cloudflare",
    "supabase", "firebase", "planetscale", "neon", "railway", "render", "fly.io",
    "lambda", "aws lambda", "azure functions", "cloud functions",
    
    # DevOps & CI/CD
    "docker", "kubernetes", "k8s", "terraform", "ansible", "puppet", "chef",
    "jenkins", "github actions", "gitlab ci", "circleci", "travis ci", "argo cd",
    "helm", "istio", "prometheus", "grafana", "datadog", "new relic", "splunk",
    "nginx", "apache", "caddy", "traefik",
    
    # Databases
    "postgresql", "postgres", "mysql", "mariadb", "mongodb", "redis", "memcached",
    "elasticsearch", "cassandra", "dynamodb", "couchdb", "neo4j", "arangodb",
    "sqlite", "sql server", "oracle", "cockroachdb", "timescaledb", "influxdb",
    "prisma", "drizzle", "typeorm", "sequelize", "sqlalchemy", "mongoose",
    
    # Version Control & Collaboration
    "git", "github", "gitlab", "bitbucket", "jira", "confluence", "notion",
    "trello", "asana", "linear", "slack",
    
    # AI/ML & Data Science
    "tensorflow", "pytorch", "keras", "scikit-learn", "sklearn", "pandas", "numpy",
    "spark", "pyspark", "hadoop", "airflow", "mlflow", "kubeflow", "dvc",
    "opencv", "nltk", "spacy", "huggingface", "transformers", "langchain",
    "openai", "gpt", "chatgpt", "llm", "large language model", "rag",
    "vector database", "pinecone", "weaviate", "milvus", "chroma", "qdrant",
    "jupyter", "notebook", "colab", "sagemaker", "vertex ai", "bedrock",
    "stable diffusion", "midjourney", "dall-e", "generative ai",
    
    # APIs & Architecture
    "graphql", "rest", "restful", "grpc", "websocket", "socket.io",
    "microservices", "serverless", "api", "oauth", "jwt", "oauth2",
    "swagger", "openapi", "postman", "insomnia",
    
    # Mobile Development
    "react native", "flutter", "ionic", "xamarin", "swiftui", "jetpack compose",
    "android", "ios", "mobile",
    
    # Testing
    "jest", "mocha", "chai", "pytest", "unittest", "junit", "testng",
    "cypress", "playwright", "selenium", "puppeteer", "testing library",
    "vitest", "enzyme", "rspec", "minitest",
    
    # Security
    "owasp", "penetration testing", "vulnerability assessment", "security audit",
    "encryption", "ssl", "tls", "https", "authentication", "authorization",
    
    # Other Tools & Concepts
    "linux", "unix", "windows server", "macos", "ubuntu", "centos", "debian",
    "agile", "scrum", "kanban", "ci/cd", "devops", "sre",
    "elasticsearch", "kibana", "logstash", "elk stack",
    "rabbitmq", "kafka", "apache kafka", "sqs", "sns", "pubsub",
    "figma", "sketch", "adobe xd", "photoshop", "illustrator"
]

# Education patterns
DEGREES = [
    "bachelor", "bachelors", "bachelor's", "b.s.", "b.s", "bs", "b.a.", "b.a", "ba",
    "b.tech", "btech", "b.e.", "be", "b.sc", "bsc", "b.com", "bcom",
    "master", "masters", "master's", "m.s.", "m.s", "ms", "m.a.", "m.a", "ma",
    "m.tech", "mtech", "m.e.", "me", "m.sc", "msc", "mba", "m.b.a",
    "ph.d", "phd", "ph.d.", "doctorate", "doctoral",
    "associate", "associates", "associate's", "a.s.", "a.a.",
    "diploma", "certificate", "certification"
]

FIELDS_OF_STUDY = [
    "computer science", "computer engineering", "software engineering",
    "information technology", "information systems", "data science",
    "electrical engineering", "electronics", "mechanical engineering",
    "mathematics", "statistics", "physics", "chemistry", "biology",
    "business administration", "management", "finance", "economics",
    "artificial intelligence", "machine learning", "cybersecurity",
    "network engineering", "cloud computing", "web development"
]


class ResumeParser:
    def __init__(self):
        self.job_titles = [title.lower() for title in JOB_TITLES]
        self.tech_skills = [skill.lower() for skill in TECH_SKILLS]
        self.degrees = [d.lower() for d in DEGREES]
        self.fields_of_study = [f.lower() for f in FIELDS_OF_STUDY]

    def extract_text_from_pdf(self, file):
        """Extract text from PDF file including tables and hyperlinks."""
        text = ""
        table_text = ""
        self.pdf_hyperlinks = []  # Store extracted hyperlinks
        
        try:
            with pdfplumber.open(file) as pdf:
                for page in pdf.pages:
                    # Extract regular text
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    
                    # Extract tables (often contains skills)
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                row_text = " ".join([str(cell) for cell in row if cell])
                                table_text += row_text + "\n"
                    
                    # Method 1: Try pdfplumber's hyperlinks property
                    try:
                        if hasattr(page, 'hyperlinks') and page.hyperlinks:
                            for link in page.hyperlinks:
                                uri = link.get('uri') or link.get('url')
                                if uri:
                                    self.pdf_hyperlinks.append(uri)
                    except Exception:
                        pass
                    
                    # Method 2: Try extracting from annotations
                    try:
                        if hasattr(page, 'annots') and page.annots:
                            for annot in page.annots:
                                # Check various annotation URI formats
                                uri = annot.get('uri') or annot.get('A', {}).get('URI')
                                if uri:
                                    self.pdf_hyperlinks.append(uri)
                    except Exception:
                        pass
                
                # Combine text and table content
                text = text + "\n" + table_text
                
                # Append hyperlinks to text so they get extracted in extract_urls
                if self.pdf_hyperlinks:
                    text += "\n" + " ".join(self.pdf_hyperlinks)
                    
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        
        # Method 3: Try PyMuPDF (fitz) as fallback for hyperlinks
        if not self.pdf_hyperlinks:
            try:
                import fitz  # PyMuPDF
                file.seek(0)  # Reset file pointer
                pdf_doc = fitz.open(stream=file.read(), filetype="pdf")
                for page in pdf_doc:
                    links = page.get_links()
                    for link in links:
                        uri = link.get('uri')
                        if uri:
                            self.pdf_hyperlinks.append(uri)
                            text += "\n" + uri
                pdf_doc.close()
            except ImportError:
                pass  # PyMuPDF not installed
            except Exception:
                pass
        
        return text

    def extract_text_from_docx(self, file):
        """Extract text from DOCX file including tables and hyperlinks."""
        try:
            doc = docx.Document(file)
            
            # Extract paragraphs
            text = "\n".join([para.text for para in doc.paragraphs])
            
            # Extract tables
            table_text = ""
            for table in doc.tables:
                for row in table.rows:
                    row_text = " ".join([cell.text for cell in row.cells if cell.text])
                    table_text += row_text + "\n"
            
            # Extract hyperlinks from document relationships
            hyperlinks = []
            try:
                # Access the document's relationship parts for hyperlinks
                for rel in doc.part.rels.values():
                    if "hyperlink" in rel.reltype:
                        if hasattr(rel, '_target') and rel._target:
                            hyperlinks.append(rel._target)
                        elif hasattr(rel, 'target_ref') and rel.target_ref:
                            hyperlinks.append(rel.target_ref)
            except Exception:
                pass  # Hyperlink extraction is optional
            
            # Append hyperlinks to text
            if hyperlinks:
                text += "\n" + " ".join(hyperlinks)
            
            return text + "\n" + table_text
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return ""

    def clean_text(self, text):
        """Remove special bullet points, extra whitespace, and formatting artifacts."""
        # Remove common bullet points and special characters
        text = re.sub(r'[\u2022\u2023\u25E6\u2043\u2219\u25CF\u25CB\u25A0\u25A1\u2610\u2611\u2612]', '', text)
        # Remove multiple spaces but preserve newlines for structure
        text = re.sub(r'[ \t]+', ' ', text)
        # Remove multiple newlines
        text = re.sub(r'\n\s*\n', '\n', text)
        return text.strip()

    def extract_name(self, text, doc):
        """Extract candidate name using NLP, heuristics, and context clues."""
        name = None
        
        # Method 1: Use Spacy NER - look for PERSON entity near the top
        if doc:
            # Only check entities in the first ~500 characters
            first_part = text[:500] if len(text) > 500 else text
            nlp = get_nlp()
            first_doc = nlp(first_part) if nlp else None
            
            if first_doc:
                for ent in first_doc.ents:
                    if ent.label_ == "PERSON":
                        candidate_name = ent.text.strip()
                        # Validate name (2-4 parts, reasonable length)
                        parts = candidate_name.split()
                        if 2 <= len(parts) <= 4 and len(candidate_name) < 50:
                            # Avoid common false positives
                            lower_name = candidate_name.lower()
                            if not any(x in lower_name for x in ['resume', 'linkedin', 'github', 'email']):
                                return candidate_name
        
        # Method 2: Look for name near contact info patterns
        email_match = re.search(r'([A-Za-z\s]+)[\s\n]*[a-zA-Z0-9._%+-]+@', text[:500])
        if email_match:
            potential_name = email_match.group(1).strip()
            parts = potential_name.split()
            if 2 <= len(parts) <= 4 and all(p[0].isupper() for p in parts if p):
                return potential_name
        
        # Method 3: First line heuristic (common in resumes)
        first_lines = text.split('\n')[:7]
        for line in first_lines:
            line = line.strip()
            if not line:
                continue
            
            parts = line.split()
            # Name is usually 2-4 words, properly capitalized
            if 2 <= len(parts) <= 4:
                # Check if it looks like a name (capitalized words)
                if all(p[0].isupper() for p in parts if p and p[0].isalpha()):
                    # Avoid lines with numbers, emails, urls, common headers
                    if not re.search(r'\d', line):
                        if '@' not in line and 'http' not in line.lower():
                            if not any(x in line.lower() for x in ['resume', 'curriculum', 'vitae', 'page', 'objective']):
                                return line
        
        # Method 4: Look for "Name:" pattern
        name_pattern = re.search(r'(?:name|full name)\s*[:\-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)', text[:500], re.IGNORECASE)
        if name_pattern:
            return name_pattern.group(1).strip()
        
        # Fallback to Spacy NER from full document
        if doc:
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    name = ent.text.strip()
                    if len(name.split()) >= 2 and len(name) < 50:
                        return name
        
        return name or "Unknown Candidate"

    def extract_email(self, text):
        """Extract email address."""
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        match = re.search(email_pattern, text)
        return match.group(0) if match else None

    def extract_phone(self, text):
        """Extract phone number with various formats."""
        patterns = [
            r'\+?1?[-.\s]?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}',
            r'\+?91[-.\s]?[0-9]{10}',  # Indian format
            r'\+?91[-.\s]?[0-9]{5}[-.\s]?[0-9]{5}',  # Indian format with space
            r'\+?[0-9]{1,3}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}[-.\s]?[0-9]{3,4}'
        ]
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0).strip()
        return None

    def extract_urls(self, text):
        """Extract GitHub, LinkedIn, and portfolio URLs."""
        urls = {
            "github": None,
            "linkedin": None,
            "portfolio": None,
            "all_urls": []
        }
        
        # Find all URLs
        url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        found_urls = re.findall(url_pattern, text)
        urls["all_urls"] = found_urls
        
        for url in found_urls:
            url_lower = url.lower()
            if "github.com" in url_lower and not urls["github"]:
                urls["github"] = url.rstrip('.,;:!?)').strip()
            elif "linkedin.com" in url_lower and not urls["linkedin"]:
                urls["linkedin"] = url.rstrip('.,;:!?)').strip()
            elif not urls["portfolio"] and "github" not in url_lower and "linkedin" not in url_lower:
                urls["portfolio"] = url.rstrip('.,;:!?)').strip()
        
        return urls

    def extract_companies(self, text, doc):
        """
        Extract EMPLOYER company names only - not tech stacks, services, or locations.
        Uses context-aware extraction: looks for companies in employment context.
        """
        companies = []
        found_companies = set()
        
        # Only educational/platform exclusions (not tech - that's handled by context)
        basic_exclusions = [
            'university', 'college', 'school', 'institute', 'academy',
            'coursera', 'udemy', 'udacity', 'edx', 'hackerrank', 'leetcode',
        ]
        
        # STEP 1: Find the Experience/Work section of the resume
        experience_section = ""
        section_patterns = [
            r'(?:work\s*experience|professional\s*experience|experience|employment|work\s*history)[\s:]*\n([\s\S]*?)(?=\n(?:education|skills|projects|certifications|achievements|awards|references|interests)|$)',
        ]
        
        for pattern in section_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                experience_section = match.group(1)
                break
        
        # If no clear section found, use full text but be stricter
        if not experience_section:
            experience_section = text
        
        # STEP 2: Extract companies from explicit employment patterns (most reliable)
        employment_patterns = [
            # "Software Engineer at Google" or "at Google as Developer"
            r'(?:at|@)\s+([A-Z][A-Za-z0-9\s\&\.\-]+?)(?:\s*[,\|\n\-–•]|\s+as\s+|\s+from\s+|\s+for\s+|\s*\()',
            
            # "Google | Software Engineer" or "Google - Senior Developer"
            r'^([A-Z][A-Za-z0-9\s\&\.\-]+?)\s*[\|–\-•]\s*(?:software|senior|junior|lead|staff|principal|engineer|developer|manager|analyst|architect|intern|consultant)',
            
            # "Worked at Google" / "Employed by Microsoft"
            r'(?:worked\s+at|working\s+at|employed\s+at|employed\s+by|joined)\s+([A-Z][A-Za-z0-9\s\&\.\-]+?)(?:\s*[,\.\n]|\s+as\s+|\s+in\s+)',
            
            # "Company: Google" or "Employer: Microsoft"
            r'(?:company|employer|organization)\s*[:\-]\s*([A-Z][A-Za-z0-9\s\&\.\-]+)',
            
            # "Google (Jan 2020 - Present)" - company followed by date
            r'^([A-Z][A-Za-z0-9\s\&\.\-]+?)\s*\(?\s*(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4})',
        ]
        
        for pattern in employment_patterns:
            matches = re.findall(pattern, experience_section, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                company = match.strip()
                company_clean = re.sub(r'\s+', ' ', company).strip(' .,-')
                
                if not company_clean or len(company_clean) < 3 or len(company_clean) > 50:
                    continue
                
                company_lower = company_clean.lower()
                
                # Skip basic exclusions
                if any(exc in company_lower for exc in basic_exclusions):
                    continue
                
                # Skip if it's just a job title
                if any(title in company_lower for title in ['engineer', 'developer', 'manager', 'analyst', 'intern', 'consultant', 'designer']):
                    continue
                
                # Skip if starts with common non-company words
                if company_lower.startswith(('the ', 'a ', 'an ', 'my ', 'our ')):
                    continue
                
                if company_clean not in found_companies:
                    found_companies.add(company_clean)
                    companies.append(company_clean)
        
        # STEP 3: Use Spacy NER only within experience section, and only for ORGs near dates
        if doc and len(companies) < 5:  # Only if we haven't found enough
            # Process just the experience section with spacy
            nlp = get_nlp()
            exp_doc = nlp(experience_section[:3000]) if nlp and experience_section else None
            
            if exp_doc:
                for ent in exp_doc.ents:
                    if ent.label_ == "ORG":
                        company = ent.text.strip()
                        company_lower = company.lower()
                        
                        # Must have at least 2 words or be a known format (Inc, Ltd, Corp, LLC)
                        word_count = len(company.split())
                        has_company_suffix = any(s in company_lower for s in ['inc', 'ltd', 'corp', 'llc', 'pvt', 'limited', 'technologies', 'solutions', 'systems', 'consulting'])
                        
                        if word_count < 2 and not has_company_suffix:
                            continue
                        
                        # Skip exclusions
                        if any(exc in company_lower for exc in basic_exclusions):
                            continue
                        
                        # Check if this ORG appears near a date (strong signal it's an employer)
                        context_start = max(0, ent.start_char - 100)
                        context_end = min(len(experience_section), ent.end_char + 100)
                        context = experience_section[context_start:context_end]
                        
                        has_date_nearby = bool(re.search(r'\b(19|20)\d{2}\b|present|current', context, re.IGNORECASE))
                        has_title_nearby = bool(re.search(r'\b(engineer|developer|manager|analyst|lead|senior|junior)\b', context, re.IGNORECASE))
                        
                        if (has_date_nearby or has_title_nearby) and company not in found_companies:
                            found_companies.add(company)
                            companies.append(company)
        
        return companies[:10]  # Limit to 10 companies

    def extract_job_titles(self, text):
        """Extract job titles from resume."""
        found_titles = []
        text_lower = text.lower()
        
        for title in self.job_titles:
            if title in text_lower:
                # Find the actual case in the text
                pattern = re.compile(re.escape(title), re.IGNORECASE)
                match = pattern.search(text)
                if match:
                    extracted = match.group(0)
                    if extracted not in found_titles:
                        found_titles.append(extracted)
        
        return found_titles[:5]  # Limit to 5 titles

    def extract_skills(self, text):
        """Extract technical skills from resume."""
        found_skills = []
        text_lower = text.lower()
        
        for skill in self.tech_skills:
            # Use word boundaries for accurate matching
            # Handle special characters in skill names like C++, C#, .NET
            escaped_skill = re.escape(skill)
            # Allow for variations like "node.js" matching "nodejs" or "node js"
            pattern = r'\b' + escaped_skill.replace(r'\.', r'\.?').replace(r'\-', r'[\-\s]?') + r'\b'
            if re.search(pattern, text_lower):
                # Normalize skill name for display
                found_skills.append(skill.title())
        
        # Remove duplicates and return
        return list(set(found_skills))

    def extract_education(self, text):
        """Extract education information including degrees and institutions."""
        education = []
        text_lower = text.lower()
        
        # Find education section
        edu_section_pattern = r'(?:education|academic|qualification|degree)[\s\S]*?(?=\n(?:experience|work|employment|skill|project|certification)|$)'
        edu_section_match = re.search(edu_section_pattern, text_lower)
        edu_text = edu_section_match.group(0) if edu_section_match else text_lower
        
        # Extract degree entries
        for degree in self.degrees:
            if degree in edu_text:
                # Try to find full context around the degree
                pattern = rf'({re.escape(degree)}[^,\n]*(?:in|of)?[^,\n]*)'
                matches = re.findall(pattern, edu_text, re.IGNORECASE)
                
                for match in matches:
                    entry = {"degree": degree.upper(), "full_text": match.strip()}
                    
                    # Try to find field of study
                    for field in self.fields_of_study:
                        if field in match.lower():
                            entry["field"] = field.title()
                            break
                    
                    # Try to find year
                    year_match = re.search(r'\b(19|20)\d{2}\b', match)
                    if year_match:
                        entry["year"] = year_match.group(0)
                    
                    if entry not in education:
                        education.append(entry)
        
        # Try to find universities/colleges using Spacy
        nlp = get_nlp()
        if nlp:
            doc = nlp(edu_text[:2000])  # Limit to first 2000 chars of edu section
            for ent in doc.ents:
                if ent.label_ == "ORG":
                    org_lower = ent.text.lower()
                    if any(x in org_lower for x in ['university', 'college', 'institute', 'school', 'academy']):
                        for edu_entry in education:
                            if 'institution' not in edu_entry:
                                edu_entry['institution'] = ent.text
                                break
                        else:
                            # If no existing entry, create one
                            if not any(e.get('institution') == ent.text for e in education):
                                education.append({"institution": ent.text})
        
        return education[:5]  # Limit to 5 education entries

    def extract_experience_dates(self, text):
        """Extract work experience dates with company associations."""
        experiences = []
        
        # Common date patterns in resumes
        date_patterns = [
            # Month Year - Month Year or Present
            r'(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{4})\s*[-–—]\s*(Present|Current|Now|(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4})',
            # MM/YYYY - MM/YYYY or Present
            r'(\d{1,2}/\d{4})\s*[-–—]\s*(Present|Current|\d{1,2}/\d{4})',
            # YYYY - YYYY or Present
            r'(\d{4})\s*[-–—]\s*(Present|Current|\d{4})'
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                experiences.append({
                    "raw_date": ' '.join(match) if isinstance(match, tuple) else match
                })
        
        return experiences

    def calculate_total_experience(self, text):
        """Calculate total years of experience from date ranges."""
        total_months = 0
        date_ranges = []
        
        # Pattern 1: Month Year - Month Year or Present
        pattern1 = r'(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{4})\s*[-–—]\s*(Present|Current|Now|(Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+(\d{4}))'
        
        matches = re.findall(pattern1, text, re.IGNORECASE)
        
        for match in matches:
            try:
                start_month = match[0]
                start_year = int(match[1])
                end_text = match[2].lower()
                
                # Parse start date
                start_date = date_parser.parse(f"{start_month} {start_year}")
                
                # Parse end date
                if end_text in ['present', 'current', 'now']:
                    end_date = datetime.now()
                else:
                    end_month = match[3] if len(match) > 3 else "Jan"
                    end_year = int(match[4]) if len(match) > 4 else start_year
                    end_date = date_parser.parse(f"{end_month} {end_year}")
                
                # Calculate duration
                if end_date > start_date:
                    diff = relativedelta(end_date, start_date)
                    months = diff.years * 12 + diff.months
                    date_ranges.append({
                        "start": start_date,
                        "end": end_date,
                        "months": months
                    })
            except Exception:
                continue
        
        # Pattern 2: YYYY - YYYY or Present (simpler fallback)
        if not date_ranges:
            pattern2 = r'(\d{4})\s*[-–—]\s*(Present|Current|\d{4})'
            matches = re.findall(pattern2, text, re.IGNORECASE)
            
            for match in matches:
                try:
                    start_year = int(match[0])
                    if match[1].lower() in ['present', 'current']:
                        end_year = datetime.now().year
                    else:
                        end_year = int(match[1])
                    
                    if end_year >= start_year and end_year <= datetime.now().year + 1:
                        months = (end_year - start_year) * 12
                        date_ranges.append({"months": months})
                except Exception:
                    continue
        
        # Sum up total experience (handle overlapping ranges later)
        total_months = sum(dr.get('months', 0) for dr in date_ranges)
        
        years = total_months // 12
        months = total_months % 12
        
        return {
            "total_years": years,
            "total_months_remainder": months,
            "total_months_raw": total_months,
            "experience_text": f"{years} years {months} months" if months else f"{years} years",
            "date_ranges_found": len(date_ranges)
        }

    def parse(self, file, file_type):
        """Main parse function to extract all entities from resume."""
        # Extract text based on file type
        if file_type == "pdf":
            text = self.extract_text_from_pdf(file)
        elif file_type == "docx":
            text = self.extract_text_from_docx(file)
        else:
            return None
        
        if not text:
            return None
        
        cleaned_text = self.clean_text(text)
        
        # Process with Spacy
        nlp = get_nlp()
        doc = nlp(cleaned_text) if nlp else None
        
        # Extract all entities
        data = {
            "name": self.extract_name(cleaned_text, doc),
            "email": self.extract_email(cleaned_text),
            "phone": self.extract_phone(cleaned_text),
            "urls": self.extract_urls(cleaned_text),
            "companies": self.extract_companies(cleaned_text, doc),
            "job_titles": self.extract_job_titles(cleaned_text),
            "skills": self.extract_skills(cleaned_text),
            "education": self.extract_education(cleaned_text),
            "experience_dates": self.extract_experience_dates(cleaned_text),
            "total_experience": self.calculate_total_experience(cleaned_text),
            "raw_text": cleaned_text
        }
        
        return data
